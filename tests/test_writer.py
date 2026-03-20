"""Tests for the DOCX writer."""

from pathlib import Path

from docx import Document

from osf_workflow import parse_osf_form, write_osf_form_to_docx


class TestWriter:
    def test_write_preserves_unmodified(self, parsed_form, osf_docx_path, tmp_path):
        out = tmp_path / "output.docx"
        write_osf_form_to_docx(parsed_form, osf_docx_path, out)
        assert out.exists()

        # Re-parse and verify structure preserved
        reparsed = parse_osf_form(out)
        assert len(reparsed.sections) == len(parsed_form.sections)
        assert len(reparsed.all_fields()) == len(parsed_form.all_fields())

    def test_write_free_text(self, parsed_form, osf_docx_path, tmp_path):
        parsed_form.edit_field("title", "Test Study Title")
        out = tmp_path / "edited.docx"
        write_osf_form_to_docx(parsed_form, osf_docx_path, out)

        doc = Document(str(out))
        fld = parsed_form.get_field("title")
        cell_text = doc.tables[fld.table_index].rows[0].cells[0].text
        assert "Test Study Title" in cell_text

    def test_write_radio(self, parsed_form, osf_docx_path, tmp_path):
        parsed_form.edit_field("study_type", "Other")
        out = tmp_path / "radio.docx"
        write_osf_form_to_docx(parsed_form, osf_docx_path, out)

        doc = Document(str(out))
        fld = parsed_form.get_field("study_type")
        if fld.table_index is not None:
            cell_text = doc.tables[fld.table_index].rows[0].cells[0].text
            assert "Other" in cell_text

    def test_write_multi_select(self, parsed_form, osf_docx_path, tmp_path):
        fld = parsed_form.get_field("blinding")
        first_opt = list(fld.response.selections.keys())[0]
        parsed_form.edit_field("blinding", {first_opt: True})
        out = tmp_path / "multiselect.docx"
        write_osf_form_to_docx(parsed_form, osf_docx_path, out)
        assert out.exists()

    def test_write_multiline_text(self, parsed_form, osf_docx_path, tmp_path):
        parsed_form.edit_field("hypotheses", "H1: Line one.\nH2: Line two.")
        out = tmp_path / "multiline.docx"
        write_osf_form_to_docx(parsed_form, osf_docx_path, out)

        doc = Document(str(out))
        fld = parsed_form.get_field("hypotheses")
        cell = doc.tables[fld.table_index].rows[0].cells[0]
        full_text = "\n".join(p.text for p in cell.paragraphs)
        assert "H1: Line one." in full_text
        assert "H2: Line two." in full_text
