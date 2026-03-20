"""Writer: OSFPreregistration → DOCX (surgical write-back).

Opens the original .docx, modifies only dirty fields, and saves to a new file.
All non-modified content remains byte-identical to the original.
"""

from __future__ import annotations

from pathlib import Path
from typing import Union

from docx import Document
from docx.oxml.ns import qn

from .schema import (
    FreeTextResponse,
    MultiSelectResponse,
    OSFPreregistration,
    RadioResponse,
    ResponseType,
)


def write_osf_form_to_docx(
    form: OSFPreregistration,
    original_docx_path: Union[str, Path],
    output_path: Union[str, Path],
) -> Path:
    """Write edited OSF form back to a .docx file.

    Only dirty fields are modified; everything else is preserved from the
    original document.

    Args:
        form: The edited OSFPreregistration object.
        original_docx_path: Path to the original (template) .docx.
        output_path: Where to save the new .docx.

    Returns:
        Path to the written file.
    """
    original_docx_path = Path(original_docx_path)
    output_path = Path(output_path)

    doc = Document(str(original_docx_path))

    dirty_count = 0
    for fld in form.all_fields():
        if not fld.dirty:
            continue
        if fld.table_index is None:
            continue

        table = doc.tables[fld.table_index]

        if fld.response_type == ResponseType.FREE_TEXT:
            if isinstance(fld.response, FreeTextResponse):
                _set_table_cell_text(table, fld.response.text)
                dirty_count += 1

        elif fld.response_type == ResponseType.RADIO:
            if isinstance(fld.response, RadioResponse) and fld.response.selected:
                _set_table_cell_text(table, fld.response.selected)
                dirty_count += 1

        elif fld.response_type == ResponseType.MULTI_SELECT:
            if isinstance(fld.response, MultiSelectResponse):
                selected = [
                    label
                    for label, checked in fld.response.selections.items()
                    if checked
                ]
                _set_table_cell_text(table, "\n".join(selected))
                dirty_count += 1

    doc.save(str(output_path))
    return output_path


def _set_table_cell_text(table, text: str) -> None:
    """Set the text content of the first cell in a 1x1 table.

    Preserves cell formatting (borders, shading). Clears existing content
    and writes new text, preserving the first paragraph's style.
    """
    cell = table.rows[0].cells[0]
    paragraphs = cell.paragraphs

    if not paragraphs:
        return

    # Set text on first paragraph, clear the rest
    first_para = paragraphs[0]

    # Handle multi-line text: split into paragraphs
    lines = text.split("\n")

    # Set first line on first paragraph
    _set_paragraph_text(first_para, lines[0])

    # Remove extra existing paragraphs (keep first)
    tc = cell._tc
    for p in paragraphs[1:]:
        tc.remove(p._p)

    # Add additional paragraphs for remaining lines
    if len(lines) > 1:
        import copy
        from lxml import etree

        insert_after = first_para._p
        for line in lines[1:]:
            new_p = etree.SubElement(tc, qn("w:p"))
            # Move it right after insert_after
            insert_after.addnext(new_p)
            # Copy paragraph properties from first para if present
            pPr = first_para._p.find(qn("w:pPr"))
            if pPr is not None:
                new_p.insert(0, copy.deepcopy(pPr))
            r = etree.SubElement(new_p, qn("w:r"))
            t = etree.SubElement(r, qn("w:t"))
            t.text = line
            t.set(qn("xml:space"), "preserve")
            insert_after = new_p


def _set_paragraph_text(paragraph, text: str) -> None:
    """Replace all text in a paragraph, preserving the first run's formatting."""
    runs = paragraph.runs

    if runs:
        # Preserve first run's formatting, clear the rest
        runs[0].text = text
        for run in runs[1:]:
            run._r.getparent().remove(run._r)
    else:
        # No runs — create one
        r = paragraph._p.makeelement(qn("w:r"), {})
        paragraph._p.append(r)
        t = r.makeelement(qn("w:t"), {})
        t.text = text
        t.set(qn("xml:space"), "preserve")
        r.append(t)
